#!/usr/bin/env python3
"""
Word Document to Markdown Converter with Text-Image Tables

This script converts a Word document to Markdown format, specifically creating
1x2 tables where:
- First row contains text content
- Second row contains the image that appears below that text

Requirements:
- python-docx
- mammoth
- markdown
- python-docx2txt
"""

import os
import sys
import subprocess
import argparse
import base64
import re
from pathlib import Path

def install_dependencies():
    """Install required Python packages if not already installed."""
    required_packages = [
        'python-docx',
        'mammoth',
        'markdown',
        'docx2txt'
    ]
    
    print("Installing required dependencies...")
    for package in required_packages:
        try:
            __import__(package.replace('-', '_'))
            print(f"✓ {package} already installed")
        except ImportError:
            print(f"Installing {package}...")
            subprocess.check_call([sys.executable, '-m', 'pip', 'install', package])
            print(f"✓ {package} installed successfully")

def main():
    """Main function to handle command line arguments and process conversion."""
    parser = argparse.ArgumentParser(description='Convert Word document to Markdown with text-image tables')
    parser.add_argument('input_file', nargs='?', default='Linux/0. Set up of LAN Segments.docx',
                       help='Input Word document file path')
    parser.add_argument('--output', '-o', help='Output Markdown file path (default: same name with .md extension)')
    parser.add_argument('--install-deps', action='store_true', help='Install required dependencies')
    parser.add_argument('--to-word', action='store_true', help='Also convert the Markdown back to Word format')
    
    args = parser.parse_args()
    
    if args.install_deps:
        install_dependencies()
        return
    
    # Check if input file exists
    input_path = Path(args.input_file)
    if not input_path.exists():
        print(f"Error: Input file '{args.input_file}' not found")
        sys.exit(1)
    
    # Determine output path
    if args.output:
        output_path = Path(args.output)
    else:
        output_path = input_path.with_suffix('.md')
    
    print(f"Converting '{input_path}' to '{output_path}'...")
    
    try:
        # Install dependencies if needed
        install_dependencies()
        
        # Perform conversion
        convert_docx_to_markdown(input_path, output_path)
        
        if args.to_word:
            word_output = output_path.with_suffix('.docx')
            convert_markdown_to_word(output_path, word_output)
            
    except Exception as e:
        print(f"Error during conversion: {e}")
        sys.exit(1)

def convert_docx_to_markdown(input_path, output_path):
    """Convert Word document to Markdown with text-image tables."""
    # Import here after ensuring dependencies are installed
    import docx
    from docx import Document
    import mammoth
    
    print("Reading Word document...")
    
    # Create images directory
    images_dir = output_path.parent / (output_path.stem + "_images")
    images_dir.mkdir(exist_ok=True)
    
    # First, try to extract with mammoth for better image handling
    try:
        with open(input_path, "rb") as docx_file:
            result = mammoth.convert_to_markdown(docx_file)
            markdown_content = result.value
            messages = result.messages
            
            if messages:
                print("Conversion messages:")
                for message in messages:
                    print(f"  {message}")
    except Exception as e:
        print(f"Mammoth conversion failed: {e}")
        print("Falling back to python-docx...")
        
        # Fallback to python-docx
        try:
            doc = Document(input_path)
            markdown_content = extract_with_python_docx(doc)
        except Exception as e2:
            print(f"python-docx conversion also failed: {e2}")
            raise
    
    # Process the markdown content to create text-image tables
    processed_content = process_content_for_tables(markdown_content, images_dir, output_path.stem)
    
    # Save the processed content
    print(f"Saving processed Markdown to '{output_path}'...")
    with open(output_path, 'w', encoding='utf-8') as f:
        f.write(processed_content)
    
    print(f"✓ Conversion completed successfully!")
    print(f"Output saved to: {output_path}")
    if images_dir.exists() and any(images_dir.iterdir()):
        print(f"Images saved to: {images_dir}")

def extract_with_python_docx(doc):
    """Extract content using python-docx as fallback."""
    content = []
    
    for paragraph in doc.paragraphs:
        if paragraph.text.strip():
            content.append(paragraph.text)
    
    # Basic conversion to markdown
    markdown_content = '\n\n'.join(content)
    return markdown_content

def process_content_for_tables(markdown_content, images_dir=None, base_name="document"):
    """Process markdown content to create 1x2 tables with text and images."""
    lines = markdown_content.split('\n')
    processed_lines = []
    image_counter = 1
    i = 0
    
    while i < len(lines):
        line = lines[i].strip()
        
        # Check if current line has text and next line(s) have an image
        if line and not line.startswith('!') and not line.startswith('|') and not line.startswith('#'):
            # Look ahead for images within the next few lines (allow for empty lines)
            image_line = None
            j = i + 1
            look_ahead_limit = 5  # Look ahead up to 5 lines
            
            # Skip empty lines and find potential image
            while j < len(lines) and (j - i) <= look_ahead_limit:
                next_line = lines[j].strip()
                if not next_line:
                    j += 1
                    continue
                elif next_line.startswith('!['):
                    image_line = next_line
                    break
                else:
                    # Found non-image content, check if it's another text line
                    # If so, stop looking for images for the current line
                    if not next_line.startswith('#') and not next_line.startswith('|'):
                        break
                    j += 1
            
            if image_line:
                # Create a 1x2 table
                table = create_markdown_table(line, image_line, images_dir, base_name, image_counter)
                processed_lines.append(table)
                processed_lines.append('')  # Add spacing after table
                image_counter += 1
                
                # Skip the processed lines up to and including the image
                i = j + 1
            else:
                # No image found, keep the text as is
                processed_lines.append(line)
                i += 1
        else:
            # Keep non-text lines as is (headers, existing tables, etc.)
            processed_lines.append(line)
            i += 1
    
    return '\n'.join(processed_lines)

def create_markdown_table(text, image, images_dir=None, base_name="document", image_counter=1):
    """Create a 1x2 markdown table with text and image."""
    # Escape pipe characters in content
    text_escaped = text.replace('|', '\\|')
    
    # Handle base64 images - extract and save as separate files if possible
    if image.startswith('![](data:image/'):
        try:
            # Extract base64 data and format
            match = re.match(r'!\[\]\(data:image/([^;]+);base64,(.+)\)', image)
            if match and images_dir:
                image_format = match.group(1)
                base64_data = match.group(2)
                
                # Decode base64 and save as file
                image_data = base64.b64decode(base64_data)
                image_filename = f"{base_name}_image_{image_counter:02d}.{image_format}"
                image_path = images_dir / image_filename
                
                with open(image_path, 'wb') as f:
                    f.write(image_data)
                
                # Create relative path for markdown
                relative_path = f"{images_dir.name}/{image_filename}"
                image_escaped = f"![Image {image_counter}]({relative_path})"
            else:
                # Fallback: use truncated version if can't save to file
                image_escaped = f"![Image {image_counter}](data:image/...)"
        except Exception as e:
            print(f"Warning: Could not extract image {image_counter}: {e}")
            image_escaped = f"![Image {image_counter}](data:image/...)"
    else:
        image_escaped = image.replace('|', '\\|')
    
    table = f"""| Content |
|---------|
| {text_escaped} |
| {image_escaped} |"""
    
    return table

def convert_markdown_to_word(markdown_path, word_path):
    """Convert Markdown back to Word format (optional feature)."""
    try:
        import markdown
        from docx import Document
        from docx.shared import Inches
        
        print(f"Converting Markdown back to Word format: '{word_path}'...")
        
        # Read markdown content
        with open(markdown_path, 'r', encoding='utf-8') as f:
            markdown_content = f.read()
        
        # Convert markdown to HTML first
        html = markdown.markdown(markdown_content, extensions=['tables'])
        
        # Create a new Word document
        doc = Document()
        
        # This is a basic implementation - would need more sophisticated parsing
        # for full HTML to Word conversion
        doc.add_heading('Converted Document', 0)
        doc.add_paragraph(f"Converted from: {markdown_path}")
        doc.add_paragraph("Note: This is a basic conversion. For full formatting, use specialized tools.")
        
        # Add the markdown content as text (basic implementation)
        for line in markdown_content.split('\n'):
            if line.strip():
                doc.add_paragraph(line)
        
        doc.save(word_path)
        print(f"✓ Word document saved to: {word_path}")
        
    except Exception as e:
        print(f"Warning: Markdown to Word conversion failed: {e}")

if __name__ == "__main__":
    main()