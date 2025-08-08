import pandas as pd
import re
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.shared import OxmlElement, qn

def add_table_borders(table):
    """Add borders to table"""
    tbl = table._tbl
    for row in tbl.tr_lst:
        for cell in row.tc_lst:
            tcPr = cell.tcPr
            tcBorders = OxmlElement('w:tcBorders')
            for border_name in ['top', 'left', 'bottom', 'right']:
                border = OxmlElement(f'w:{border_name}')
                border.set(qn('w:val'), 'single')
                border.set(qn('w:sz'), '4')
                border.set(qn('w:space'), '0')
                border.set(qn('w:color'), '000000')
                tcBorders.append(border)
            tcPr.append(tcBorders)

def extract_aspects():
    """Extract all aspects from the Excel file"""
    df = pd.read_excel('Linux Marking Scheme.xlsx')
    aspects = []
    
    for i in range(len(df)):
        aspect_desc = df.iloc[i, 4]  # Column 4: Aspect Description
        commands = df.iloc[i, 6]     # Column 6: Commands
        expected_output = df.iloc[i, 7]  # Column 7: Expected Output
        
        # Skip if aspect description is empty or NaN
        if pd.isna(aspect_desc) or not isinstance(aspect_desc, str):
            continue
        
        # Skip if commands is empty or NaN
        if pd.isna(commands) or not isinstance(commands, str):
            continue
        
        # Clean the commands - remove the grading wrapper and "Executed command on..." lines
        clean_commands = commands
        
        # Remove lines starting with "./grading -v -t"
        clean_commands = re.sub(r'./grading -v -t [^\n]*\n?', '', clean_commands)
        
        # Remove lines starting with "Executed command on"
        clean_commands = re.sub(r'Executed command on [^\n]*=>\n?', '', clean_commands)
        
        # Clean up extra whitespace and newlines
        clean_commands = re.sub(r'\n\s*\n', '\n', clean_commands)
        clean_commands = clean_commands.strip()
        
        # Skip if no commands remain after cleaning
        if not clean_commands:
            continue
        
        # Format expected output
        clean_expected = str(expected_output) if pd.notna(expected_output) else ""
        
        aspects.append({
            'aspect': aspect_desc,
            'commands': clean_commands,
            'expected_output': clean_expected
        })
    
    return aspects

def create_word_document():
    """Create Word document with all aspects"""
    # Create new document
    doc = Document()
    
    # Set document title
    title = doc.add_heading('Linux Network Systems Administration - Commands and Expected Outputs', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add spacing
    doc.add_paragraph()
    
    # Extract aspects
    aspects = extract_aspects()
    
    # Create table for each aspect
    for i, aspect in enumerate(aspects, 1):
        # Add aspect heading
        heading = doc.add_heading(f'{i}. {aspect["aspect"]}', level=1)
        
        # Create table with 2 rows and 1 column
        table = doc.add_table(rows=2, cols=1)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = True
        
        # Set table width
        table.columns[0].width = Inches(6.5)
        
        # Command row
        cmd_cell = table.cell(0, 0)
        cmd_cell.text = "Command"
        cmd_para = cmd_cell.paragraphs[0]
        cmd_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Make command header bold
        run = cmd_para.runs[0]
        run.bold = True
        run.font.size = Pt(12)
        
        # Add command content
        cmd_content = cmd_cell.add_paragraph()
        cmd_content.text = aspect['commands']
        cmd_content.style = 'Normal'
        
        # Expected Output row
        output_cell = table.cell(1, 0)
        output_cell.text = "Expected Output"
        output_para = output_cell.paragraphs[0]
        output_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Make expected output header bold
        run = output_para.runs[0]
        run.bold = True
        run.font.size = Pt(12)
        
        # Add expected output content
        output_content = output_cell.add_paragraph()
        output_content.text = aspect['expected_output']
        output_content.style = 'Normal'
        
        # Add borders to table
        add_table_borders(table)
        
        # Add spacing between aspects
        doc.add_paragraph()
    
    # Save document
    doc.save('Linux_Commands_and_Expected_Outputs.docx')
    print(f"Document created successfully with {len(aspects)} aspects!")
    print("Saved as: Linux_Commands_and_Expected_Outputs.docx")

if __name__ == "__main__":
    create_word_document() 